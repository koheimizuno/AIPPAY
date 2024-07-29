from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Inches, Cm, Mm
from docx.oxml.ns import qn
import io
from datetime import datetime, timezone

class DocxMaker(object):
    """
    Microsoft Word形式のファイルを生成する
    ※python-docxのラッパー
    """

    def __init__(self, default_font='游明朝', margin=None):
        """
        コンストラクター
        """
        self.__doc = Document()
        self.__doc.core_properties.created = datetime.now(timezone.utc)
        self.__doc.core_properties.modified = datetime.now(timezone.utc)
        self.__margin = margin
        self.__init_section()
        if not default_font is None and default_font != '':
            self.__default_font = default_font
        else:
            self.__default_font = "游明朝"

    def __init_section(self):
        """
        セクションの初期化
        """
        if not self.__margin is None:
            self.current.top_margin = Mm(self.__margin[0])
            self.current.right_margin = Mm(self.__margin[1])
            self.current.bottom_margin = Mm(self.__margin[2])
            self.current.left_margin = Mm(self.__margin[3])
        else:
            self.current.top_margin = Mm(16)
            self.current.right_margin = Mm(20)
            self.current.bottom_margin = Mm(18.5)
            self.current.left_margin = Mm(22)

    def new_page(self):
        """
        改ページ
        """
        self.document.add_section(WD_SECTION.NEW_PAGE)
        self.__init_section()

    @property
    def document(self):
        return self.__doc

    @property
    def current(self):
        return self.document.sections[-1]
    
    @property
    def header(self):
        return self.current.header

    def add_heading(self, text, font_size=10.5, center=False, level=0):
        """
        見出しの追加
        """
        p = self.document.add_heading("", level=level)
        run = p.add_run(text)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(font_size)
        set_run_font(run, self.__default_font)
        p.paragraph_format.line_spacing = Pt(font_size)
        p.paragraph_format.space_after = Pt(0)

    def add_paragraph(self, text, font_size=10.5, indent=0.0, center=False, right=False, follow_indent=0.0):
        """
        段落の追加
        """
        p = self.document.add_paragraph("")
        run = p.add_run(text)
        if right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(font_size)
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent)
        if follow_indent > 0:
            p.paragraph_format.left_indent = Cm(follow_indent)
            p.paragraph_format.first_line_indent = Cm(indent - follow_indent)
        set_run_font(run, self.__default_font)
        #p.paragraph_format.line_spacing = Pt(font_size)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)

    def add_paragraph_into_header(self, text, font_size=10.5, indent=0.0, center=False, right=False):
        """
        段落の追加
        """
        p = self.document.sections[0].header.add_paragraph("")
        run = p.add_run(text)
        if right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(font_size)
        set_run_font(run, self.__default_font)
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent)
        #p.paragraph_format.line_spacing = Pt(font_size)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)

    def add_picture(self, image, width=None, center=False, right=False):
        """
        画像の追加
        """
        if width:
            p = self.document.add_picture(image, width=Pt(width))
        else:
            p = self.document.add_picture(image)
        p = self.document.paragraphs[-1]
        if right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table(self, values, font_size=10.5, underline=None):
        """
        表の追加
        """
        table = self.document.add_table(rows=len(values), cols=max([len(x) for x in values]))
        #table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        for r in range(len(values)):
            for c in range(len(values[r])):
                cell = table.cell(r, c)
                p = cell.paragraphs[0]
                run = p.add_run(str(values[r][c]))
                run.font.size = Pt(font_size)
                set_run_font(run, self.__default_font)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(0)
                if c == 0:
                    cell.width = Cm(4)
                elif c == 1:
                    cell.width = Cm(12)
                if not underline is None and underline[r][c]:
                    run.underline = True

    #def save(self, path):
    #    """
    #    ファイルに保存
    #    """
    #    self.document.save(path)

    def get_binary(self):
        """
        バイナリーを取得
        """
        with io.BytesIO() as buff:
            self.document.save(buff)
            return buff.getvalue()

def set_run_font(run, font_name):
    """
    フォントの設定（日本語対応）
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

if __name__ == '__main__':
    d = DocxMaker()
    d.add_paragraph('タイトル', center=True)
    bin = d.get_binary()
    with open('out.docx', 'wb') as fout:
        fout.write(bin)
